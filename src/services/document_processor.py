import os
import hashlib
import chardet
from typing import Optional, List, Dict, Any, BinaryIO
from pathlib import Path
import logging

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import markdown

from src.config.settings import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Handles document text extraction from various formats.
    """

    SUPPORTED_FORMATS = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'txt': ['text/plain'],
        'md': ['text/markdown', 'text/x-markdown'],
    }

    @staticmethod
    def calculate_file_hash(file_content: bytes) -> str:
        """Calculate SHA-256 hash of file content."""
        return hashlib.sha256(file_content).hexdigest()

    @staticmethod
    def detect_encoding(file_content: bytes) -> str:
        """Detect file encoding for text files."""
        result = chardet.detect(file_content)
        return result['encoding'] or 'utf-8'

    def extract_text(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Extract text from a document based on its type.

        Returns:
            Dict containing:
            - text: extracted text content
            - pages: list of page texts (for PDFs)
            - total_pages: number of pages
            - metadata: additional document metadata
        """
        file_type = file_type.lower()

        if file_type == 'pdf':
            return self._extract_pdf_text(file_path)
        elif file_type == 'docx':
            return self._extract_docx_text(file_path)
        elif file_type == 'txt':
            return self._extract_txt_text(file_path)
        elif file_type == 'md':
            return self._extract_md_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using pdfplumber for better accuracy."""
        pages = []
        metadata = {}

        try:
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                if pdf.metadata:
                    metadata = {
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                        'producer': pdf.metadata.get('Producer', ''),
                        'creation_date': str(pdf.metadata.get('CreationDate', '')),
                        'modification_date': str(pdf.metadata.get('ModDate', '')),
                    }

                # Extract text from each page
                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text() or ''
                        pages.append({
                            'page_number': i + 1,
                            'text': text,
                            'char_count': len(text),
                        })
                    except Exception as e:
                        logger.warning(f"Error extracting page {i+1}: {e}")
                        pages.append({
                            'page_number': i + 1,
                            'text': '',
                            'char_count': 0,
                        })

                # Combine all text
                full_text = '\n\n'.join([p['text'] for p in pages])

                return {
                    'text': full_text,
                    'pages': pages,
                    'total_pages': len(pages),
                    'metadata': metadata,
                }

        except Exception as e:
            logger.error(f"PDF extraction failed with pdfplumber, trying PyPDF2: {e}")
            return self._extract_pdf_text_fallback(file_path)

    def _extract_pdf_text_fallback(self, file_path: str) -> Dict[str, Any]:
        """Fallback PDF extraction using PyPDF2."""
        pages = []
        metadata = {}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    }

                # Extract text from each page
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        pages.append({
                            'page_number': i + 1,
                            'text': text,
                            'char_count': len(text),
                        })
                    except Exception as e:
                        logger.warning(f"Error extracting page {i+1} with PyPDF2: {e}")
                        pages.append({
                            'page_number': i + 1,
                            'text': '',
                            'char_count': 0,
                        })

                # Combine all text
                full_text = '\n\n'.join([p['text'] for p in pages])

                return {
                    'text': full_text,
                    'pages': pages,
                    'total_pages': len(pages),
                    'metadata': metadata,
                }

        except Exception as e:
            logger.error(f"PDF extraction failed completely: {e}")
            raise

    def _extract_docx_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)

            # Extract metadata
            metadata = {}
            if doc.core_properties:
                metadata = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                    'revision': doc.core_properties.revision or 0,
                }

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    table_text.append('\t'.join(row_text))
                tables_text.append('\n'.join(table_text))

            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\n' + '\n\n'.join(tables_text)

            return {
                'text': full_text,
                'pages': [{'page_number': 1, 'text': full_text, 'char_count': len(full_text)}],
                'total_pages': 1,
                'metadata': metadata,
            }

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_txt_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file with encoding detection."""
        try:
            # Read file content
            with open(file_path, 'rb') as file:
                raw_content = file.read()

            # Detect encoding
            encoding = self.detect_encoding(raw_content)

            # Decode text
            text = raw_content.decode(encoding, errors='replace')

            return {
                'text': text,
                'pages': [{'page_number': 1, 'text': text, 'char_count': len(text)}],
                'total_pages': 1,
                'metadata': {'encoding': encoding},
            }

        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise

    def _extract_md_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Markdown file."""
        try:
            # Read file content
            with open(file_path, 'rb') as file:
                raw_content = file.read()

            # Detect encoding
            encoding = self.detect_encoding(raw_content)

            # Decode text
            text = raw_content.decode(encoding, errors='replace')

            # Convert markdown to plain text (optional, keeping markdown for now)
            # html = markdown.markdown(text)
            # We'll keep the markdown format as it's useful for chunking

            return {
                'text': text,
                'pages': [{'page_number': 1, 'text': text, 'char_count': len(text)}],
                'total_pages': 1,
                'metadata': {'encoding': encoding, 'format': 'markdown'},
            }

        except Exception as e:
            logger.error(f"Markdown extraction failed: {e}")
            raise

    def validate_file(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Validate file before processing.

        Returns:
            Dict containing:
            - valid: bool
            - file_type: detected file type
            - file_hash: SHA-256 hash
            - error: error message if invalid
        """
        # Check file size
        file_size = len(file_content)
        if file_size > settings.max_upload_size_bytes:
            return {
                'valid': False,
                'error': f'File size exceeds maximum of {settings.max_upload_size_mb}MB',
            }

        # Check file extension
        file_ext = Path(file_name).suffix.lower().lstrip('.')
        if file_ext not in settings.allowed_file_types:
            return {
                'valid': False,
                'error': f'File type .{file_ext} not allowed. Allowed types: {", ".join(settings.allowed_file_types)}',
            }

        # Calculate file hash
        file_hash = self.calculate_file_hash(file_content)

        # Basic file signature validation
        if not self._validate_file_signature(file_content, file_ext):
            return {
                'valid': False,
                'error': f'File content does not match expected format for .{file_ext}',
            }

        return {
            'valid': True,
            'file_type': file_ext,
            'file_hash': file_hash,
            'file_size': file_size,
        }

    def _validate_file_signature(self, file_content: bytes, file_ext: str) -> bool:
        """Validate file signature based on magic bytes."""
        if not file_content:
            return False

        # File signatures (magic bytes)
        signatures = {
            'pdf': b'%PDF',
            'docx': b'PK\x03\x04',  # ZIP archive (DOCX is a ZIP)
            'txt': None,  # Text files don't have a signature
            'md': None,   # Markdown files don't have a signature
        }

        expected_signature = signatures.get(file_ext)

        if expected_signature is None:
            # No signature check needed
            return True

        return file_content.startswith(expected_signature)