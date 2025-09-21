"""
Sample files and data for testing.
"""

import os
import tempfile
from typing import Dict, Any
import pytest


class TestFileGenerator:
    """Generate test files for various document types."""

    @staticmethod
    def create_sample_pdf(content: str = "Sample PDF content") -> bytes:
        """Create a minimal PDF file for testing."""
        # Basic PDF structure
        pdf_content = f"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length {len(content)}
>>
stream
BT
/F1 12 Tf
72 720 Td
({content}) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000074 00000 n
0000000120 00000 n
0000000179 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
{300 + len(content)}
%%EOF"""
        return pdf_content.encode('utf-8')

    @staticmethod
    def create_sample_docx(content: str = "Sample DOCX content") -> bytes:
        """Create a minimal DOCX file for testing."""
        from docx import Document

        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph(content)

        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def create_corrupted_file(file_type: str) -> bytes:
        """Create a corrupted file for testing error handling."""
        if file_type == "pdf":
            return b"Not a PDF file but has PDF extension"
        elif file_type == "docx":
            return b"Not a DOCX file but has DOCX extension"
        else:
            return b"Corrupted file content"

    @staticmethod
    def create_large_text_file(size_kb: int = 100) -> str:
        """Create large text content for performance testing."""
        base_text = "This is a test sentence for performance testing. It contains multiple words and punctuation marks. "
        repeat_count = (size_kb * 1024) // len(base_text)
        return base_text * repeat_count

    @staticmethod
    def create_multilingual_text() -> str:
        """Create text with multiple languages for encoding testing."""
        return """
English: Hello, world! This is a test document.
Español: ¡Hola, mundo! Este es un documento de prueba.
Français: Bonjour, le monde! Ceci est un document de test.
Deutsch: Hallo, Welt! Dies ist ein Testdokument.
中文: 你好，世界！这是一个测试文档。
العربية: مرحبا بالعالم! هذا مستند اختبار.
Русский: Привет, мир! Это тестовый документ.
日本語: こんにちは、世界！これはテストドキュメントです。
한국어: 안녕하세요, 세계! 이것은 테스트 문서입니다.
हिन्दी: नमस्ते, दुनिया! यह एक परीक्षण दस्तावेज़ है।
"""

    @staticmethod
    def create_structured_markdown() -> str:
        """Create structured markdown content for testing."""
        return """# Main Document Title

## Introduction

This is a **test document** with various markdown elements.

### Features

- Bullet point 1
- Bullet point 2
- Bullet point 3

#### Sub-features

1. Numbered item 1
2. Numbered item 2
3. Numbered item 3

## Code Examples

Here's some Python code:

```python
def hello_world():
    print("Hello, world!")
    return True
```

And some JavaScript:

```javascript
function helloWorld() {
    console.log("Hello, world!");
    return true;
}
```

## Tables

| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Value 1  | Value 2  | Value 3  |
| Data A   | Data B   | Data C   |

## Links and Images

[Example Link](https://example.com)

> This is a blockquote
> with multiple lines

## Conclusion

This document demonstrates various markdown formatting options.
"""

    @staticmethod
    def create_technical_content() -> str:
        """Create technical content for testing domain-specific processing."""
        return """
# Technical Specification Document

## API Endpoints

### POST /api/v1/documents
Create a new document resource.

**Request Body:**
```json
{
    "title": "string",
    "content": "string",
    "metadata": {
        "author": "string",
        "created_at": "2023-01-01T00:00:00Z"
    }
}
```

**Response:**
```json
{
    "id": "uuid",
    "status": "created",
    "document_url": "https://api.example.com/documents/uuid"
}
```

### Mathematical Formulas

The quadratic formula: x = (-b ± √(b²-4ac)) / 2a

Einstein's mass-energy equivalence: E = mc²

## Configuration Examples

Environment variables:
- `DATABASE_URL`: PostgreSQL connection string
- `REDIS_URL`: Redis connection string
- `API_KEY`: Authentication token

Docker command:
```bash
docker run -p 8000:8000 -e DATABASE_URL=postgresql://user:pass@host:5432/db myapp
```

## Error Codes

| Code | Description | Resolution |
|------|-------------|------------|
| 400  | Bad Request | Check request syntax |
| 401  | Unauthorized | Provide valid authentication |
| 429  | Rate Limited | Reduce request frequency |
| 500  | Server Error | Contact support |
"""


@pytest.fixture
def test_file_generator():
    """Provide test file generator."""
    return TestFileGenerator()


@pytest.fixture
def sample_files_dict(temp_dir, test_file_generator):
    """Create dictionary of sample files for testing."""
    files = {}

    # Text files
    files['simple.txt'] = os.path.join(temp_dir, 'simple.txt')
    with open(files['simple.txt'], 'w', encoding='utf-8') as f:
        f.write("Simple text content for testing.")

    files['multilingual.txt'] = os.path.join(temp_dir, 'multilingual.txt')
    with open(files['multilingual.txt'], 'w', encoding='utf-8') as f:
        f.write(test_file_generator.create_multilingual_text())

    files['large.txt'] = os.path.join(temp_dir, 'large.txt')
    with open(files['large.txt'], 'w', encoding='utf-8') as f:
        f.write(test_file_generator.create_large_text_file(50))  # 50KB

    # Markdown files
    files['structured.md'] = os.path.join(temp_dir, 'structured.md')
    with open(files['structured.md'], 'w', encoding='utf-8') as f:
        f.write(test_file_generator.create_structured_markdown())

    files['technical.md'] = os.path.join(temp_dir, 'technical.md')
    with open(files['technical.md'], 'w', encoding='utf-8') as f:
        f.write(test_file_generator.create_technical_content())

    # PDF files
    files['sample.pdf'] = os.path.join(temp_dir, 'sample.pdf')
    with open(files['sample.pdf'], 'wb') as f:
        f.write(test_file_generator.create_sample_pdf())

    files['corrupted.pdf'] = os.path.join(temp_dir, 'corrupted.pdf')
    with open(files['corrupted.pdf'], 'wb') as f:
        f.write(test_file_generator.create_corrupted_file('pdf'))

    # DOCX files
    files['sample.docx'] = os.path.join(temp_dir, 'sample.docx')
    with open(files['sample.docx'], 'wb') as f:
        f.write(test_file_generator.create_sample_docx())

    files['corrupted.docx'] = os.path.join(temp_dir, 'corrupted.docx')
    with open(files['corrupted.docx'], 'wb') as f:
        f.write(test_file_generator.create_corrupted_file('docx'))

    return files


@pytest.fixture
def test_embeddings_data():
    """Provide test embeddings data."""
    return {
        "dimension": 1536,
        "sample_embeddings": [
            [0.1, -0.2, 0.3] * 512,  # 1536 dimensions
            [-0.1, 0.2, -0.3] * 512,
            [0.05, -0.1, 0.15] * 512,
        ],
        "similarity_pairs": [
            ([1.0, 0.0, 0.0], [1.0, 0.0, 0.0], 1.0),  # Identical
            ([1.0, 0.0, 0.0], [0.0, 1.0, 0.0], 0.0),  # Orthogonal
            ([1.0, 0.0, 0.0], [-1.0, 0.0, 0.0], 0.0), # Opposite
            ([1.0, 1.0, 0.0], [1.0, 1.0, 0.0], 1.0),  # Identical
        ]
    }


@pytest.fixture
def chunking_test_cases():
    """Provide test cases for text chunking."""
    return [
        {
            "name": "simple_sentences",
            "text": "First sentence. Second sentence. Third sentence.",
            "expected_chunks": 1,
            "chunk_size_max": 100,
        },
        {
            "name": "paragraph_breaks",
            "text": "First paragraph.\n\nSecond paragraph.\n\nThird paragraph.",
            "expected_chunks": 3,
            "chunk_size_max": 50,
        },
        {
            "name": "long_text",
            "text": "This is a very long text. " * 100,
            "expected_chunks": 10,
            "chunk_size_max": 200,
        },
        {
            "name": "single_long_sentence",
            "text": "This is one extremely long sentence that should be broken down into smaller chunks because it exceeds the maximum chunk size limit that we have set for testing purposes.",
            "expected_chunks": 2,
            "chunk_size_max": 50,
        }
    ]


@pytest.fixture
def api_test_cases():
    """Provide test cases for API testing."""
    return {
        "valid_uploads": [
            {
                "filename": "test.txt",
                "content": b"Valid text content",
                "content_type": "text/plain",
                "expected_status": 200
            },
            {
                "filename": "test.md",
                "content": b"# Valid markdown content",
                "content_type": "text/markdown",
                "expected_status": 200
            }
        ],
        "invalid_uploads": [
            {
                "filename": "test.exe",
                "content": b"Executable content",
                "content_type": "application/x-executable",
                "expected_status": 400
            },
            {
                "filename": "test.txt",
                "content": b"x" * (100 * 1024 * 1024),  # 100MB
                "content_type": "text/plain",
                "expected_status": 400
            }
        ],
        "malicious_uploads": [
            {
                "filename": "../../../etc/passwd",
                "content": b"root:x:0:0:root:/root:/bin/bash",
                "content_type": "text/plain",
                "expected_status": 400
            },
            {
                "filename": "script.txt",
                "content": b"<script>alert('xss')</script>",
                "content_type": "text/plain",
                "expected_status": 200  # Should accept but sanitize
            }
        ]
    }


@pytest.fixture
def performance_test_data():
    """Provide data for performance testing."""
    return {
        "file_sizes": [1024, 10240, 102400, 1024000],  # 1KB to 1MB
        "batch_sizes": [1, 10, 50, 100],
        "concurrent_users": [1, 5, 10, 20],
        "text_lengths": [100, 1000, 10000, 50000],  # characters
    }


@pytest.fixture
def security_test_data():
    """Provide data for security testing."""
    return {
        "sql_injection_attempts": [
            "'; DROP TABLE documents; --",
            "' UNION SELECT * FROM profiles --",
            "' OR '1'='1",
            "admin'/*"
        ],
        "xss_attempts": [
            "<script>alert('xss')</script>",
            "javascript:alert('xss')",
            "<img src=x onerror=alert('xss')>",
            "<svg onload=alert('xss')>"
        ],
        "path_traversal_attempts": [
            "../../../etc/passwd",
            "..\\..\\..\\windows\\system32\\hosts",
            "/etc/passwd",
            "C:\\windows\\system32\\config\\sam"
        ],
        "malicious_file_signatures": [
            b"MZ\x90\x00",  # PE executable
            b"\x7fELF",     # ELF executable
            b"#!/bin/sh",   # Shell script
            b"PK\x03\x04",  # ZIP (could be malicious DOCX)
        ]
    }