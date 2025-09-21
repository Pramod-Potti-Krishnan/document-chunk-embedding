"""
Global test configuration and fixtures.
"""

import os
import tempfile
import asyncio
import pytest
from typing import AsyncGenerator, Generator, Dict, Any
from unittest.mock import Mock, AsyncMock, patch
from uuid import uuid4
import hashlib

# FastAPI testing
from fastapi.testclient import TestClient
from httpx import AsyncClient

# Database testing
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
from sqlalchemy.pool import StaticPool

# Test environment setup
os.environ.update({
    "ENVIRONMENT": "test",
    "TESTING": "True",
    "LOG_LEVEL": "DEBUG",
    "DATABASE_URL": "sqlite:///./test.db",
    "SUPABASE_URL": "http://localhost:54321",
    "SUPABASE_KEY": "test_key",
    "OPENAI_API_KEY": "test_openai_key",
    "REDIS_HOST": "localhost",
    "REDIS_PORT": "6379",
    "REDIS_DB": "15",  # Use different DB for tests
    "JWT_SECRET_KEY": "test_secret_key_for_testing_only",
    "CORS_ORIGINS": "http://localhost:3000,http://localhost:8080",
    "MAX_UPLOAD_SIZE_MB": "50",
    "CHUNK_SIZE_MIN": "100",
    "CHUNK_SIZE_MAX": "1000",
    "CHUNK_OVERLAP": "50",
})

# Import after environment setup
from src.main import app
from src.models.database import Base, Document, DocumentChunk, Profile
from src.core.database import get_db
from src.config.settings import settings


# Pytest configuration
def pytest_configure(config):
    """Configure pytest with custom markers."""
    config.addinivalue_line("markers", "unit: Unit tests")
    config.addinivalue_line("markers", "integration: Integration tests")
    config.addinivalue_line("markers", "e2e: End-to-end tests")
    config.addinivalue_line("markers", "performance: Performance tests")
    config.addinivalue_line("markers", "security: Security tests")


# Event loop fixture for async tests
@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.new_event_loop()
    yield loop
    loop.close()


# Database fixtures
@pytest.fixture(scope="session")
def test_db_engine():
    """Create test database engine."""
    engine = create_engine(
        "sqlite:///./test.db",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    yield engine
    Base.metadata.drop_all(bind=engine)


@pytest.fixture(scope="function")
def test_db_session(test_db_engine) -> Generator[Session, None, None]:
    """Create test database session."""
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_db_engine)
    session = TestingSessionLocal()

    try:
        yield session
    finally:
        session.rollback()
        session.close()


@pytest.fixture(scope="function")
def override_get_db(test_db_session):
    """Override get_db dependency for testing."""
    def _override_get_db():
        try:
            yield test_db_session
        finally:
            pass

    app.dependency_overrides[get_db] = _override_get_db
    yield
    app.dependency_overrides.clear()


# FastAPI client fixtures
@pytest.fixture(scope="function")
def client(override_get_db) -> TestClient:
    """Create test client."""
    return TestClient(app)


@pytest.fixture(scope="function")
async def async_client(override_get_db) -> AsyncGenerator[AsyncClient, None]:
    """Create async test client."""
    async with AsyncClient(app=app, base_url="http://test") as ac:
        yield ac


# Authentication fixtures
@pytest.fixture
def test_user() -> Dict[str, Any]:
    """Create test user data."""
    return {
        "user_id": str(uuid4()),
        "email": "test@example.com",
        "session_id": str(uuid4()),
        "project_id": str(uuid4()),
    }


@pytest.fixture
def auth_headers(test_user) -> Dict[str, str]:
    """Create authentication headers for test user."""
    from src.core.auth import create_access_token
    token = create_access_token({"sub": test_user["user_id"]})
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def mock_current_user(test_user):
    """Mock current user dependency."""
    with patch("src.core.auth.get_current_user", return_value=test_user):
        yield test_user


# File fixtures
@pytest.fixture
def temp_dir() -> Generator[str, None, None]:
    """Create temporary directory for test files."""
    with tempfile.TemporaryDirectory() as temp_dir:
        yield temp_dir


@pytest.fixture
def sample_text_content() -> str:
    """Sample text content for testing."""
    return """# Sample Document

This is a sample document for testing purposes.

## Chapter 1: Introduction

Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.

## Chapter 2: Content

Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.

Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo.

## Conclusion

Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt.
"""


@pytest.fixture
def sample_pdf_file(temp_dir) -> str:
    """Create a sample PDF file for testing."""
    # This would normally use a library like reportlab to create a PDF
    # For now, we'll create a mock PDF with basic header
    pdf_content = b"%PDF-1.4\nSample PDF content for testing"
    pdf_path = os.path.join(temp_dir, "sample.pdf")
    with open(pdf_path, "wb") as f:
        f.write(pdf_content)
    return pdf_path


@pytest.fixture
def sample_txt_file(temp_dir, sample_text_content) -> str:
    """Create a sample text file for testing."""
    txt_path = os.path.join(temp_dir, "sample.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(sample_text_content)
    return txt_path


@pytest.fixture
def sample_md_file(temp_dir, sample_text_content) -> str:
    """Create a sample markdown file for testing."""
    md_path = os.path.join(temp_dir, "sample.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(sample_text_content)
    return md_path


@pytest.fixture
def sample_docx_file(temp_dir) -> str:
    """Create a sample DOCX file for testing."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Sample Document", 0)
    doc.add_paragraph("This is a sample DOCX document for testing.")
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit.")

    docx_path = os.path.join(temp_dir, "sample.docx")
    doc.save(docx_path)
    return docx_path


# Mock service fixtures
@pytest.fixture
def mock_embeddings_service():
    """Mock embeddings service."""
    from src.services.embeddings_service import MockEmbeddingsService
    return MockEmbeddingsService()


@pytest.fixture
def mock_redis_client():
    """Mock Redis client."""
    mock_redis = Mock()
    mock_redis.ping.return_value = True
    mock_redis.get.return_value = None
    mock_redis.set.return_value = True
    mock_redis.delete.return_value = 1
    return mock_redis


@pytest.fixture
def mock_supabase_client():
    """Mock Supabase client."""
    mock_supabase = Mock()
    mock_supabase.auth.get_session.return_value = {"access_token": "test_token"}
    mock_supabase.storage.from_.return_value.upload.return_value = {"path": "test/path"}
    return mock_supabase


@pytest.fixture
def mock_celery_task():
    """Mock Celery task."""
    mock_task = Mock()
    mock_task.delay.return_value.id = str(uuid4())
    mock_task.delay.return_value.status = "PENDING"
    return mock_task


# Database test data fixtures
@pytest.fixture
def sample_document(test_db_session, test_user) -> Document:
    """Create a sample document in the database."""
    document = Document(
        id=uuid4(),
        user_id=test_user["user_id"],
        session_id=test_user["session_id"],
        project_id=test_user["project_id"],
        filename="sample.txt",
        file_type="txt",
        file_size_bytes=1024,
        file_hash=hashlib.sha256(b"sample content").hexdigest(),
        mime_type="text/plain",
        status="completed",
        storage_path="/tmp/sample.txt",
        metadata={"test": True},
        tags=["test", "sample"]
    )
    test_db_session.add(document)
    test_db_session.commit()
    test_db_session.refresh(document)
    return document


@pytest.fixture
def sample_chunks(test_db_session, sample_document) -> list:
    """Create sample document chunks in the database."""
    chunks = []
    for i in range(3):
        chunk = DocumentChunk(
            id=uuid4(),
            document_id=sample_document.id,
            chunk_index=i,
            text_content=f"This is chunk {i} of the document.",
            chunk_size=50 + i * 10,
            token_count=15 + i * 3,
            page_number=1,
            start_char=i * 60,
            end_char=(i + 1) * 60,
            embedding=[0.1] * 1536,  # Mock embedding
            metadata={"chunk_type": "paragraph"}
        )
        chunks.append(chunk)
        test_db_session.add(chunk)

    test_db_session.commit()
    for chunk in chunks:
        test_db_session.refresh(chunk)
    return chunks


@pytest.fixture
def sample_profile(test_db_session, test_user) -> Profile:
    """Create a sample user profile in the database."""
    profile = Profile(
        user_id=test_user["user_id"],
        email=test_user["email"],
        storage_used_mb=10.5,
        storage_limit_mb=1000.0,
        documents_count=5,
        plan_type="free",
        api_calls_count=100,
        api_calls_limit=1000
    )
    test_db_session.add(profile)
    test_db_session.commit()
    test_db_session.refresh(profile)
    return profile


# Utility fixtures
@pytest.fixture
def file_hash():
    """Utility function to calculate file hash."""
    def _calculate_hash(content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()
    return _calculate_hash


@pytest.fixture
def upload_file_data():
    """Create upload file data for testing."""
    def _create_upload_data(filename: str, content: bytes, content_type: str = "text/plain"):
        return {
            "file": (filename, content, content_type),
            "user_id": str(uuid4()),
            "session_id": str(uuid4()),
            "project_id": "-",
            "metadata": "{}",
            "tags": "[]"
        }
    return _create_upload_data


# Performance testing fixtures
@pytest.fixture
def large_text_content() -> str:
    """Generate large text content for performance testing."""
    base_text = "This is a test sentence for performance testing. " * 100
    return base_text * 1000  # ~50KB of text


@pytest.fixture
def large_file_content() -> bytes:
    """Generate large file content for testing."""
    return b"Test content for large file testing. " * 50000  # ~1.5MB


# Security testing fixtures
@pytest.fixture
def malicious_file_content() -> bytes:
    """Create malicious file content for security testing."""
    return b"<script>alert('xss')</script>" + b"A" * 1000


@pytest.fixture
def invalid_token():
    """Create invalid JWT token for testing."""
    return "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.invalid.token"


# Cleanup fixtures
@pytest.fixture(autouse=True)
def cleanup_temp_files():
    """Automatically cleanup temporary files after each test."""
    yield
    # Cleanup logic can be added here if needed
    pass


# API response validation fixtures
@pytest.fixture
def validate_response():
    """Utility to validate API responses."""
    def _validate(response, expected_status: int = 200, expected_keys: list = None):
        assert response.status_code == expected_status
        if expected_keys:
            data = response.json()
            for key in expected_keys:
                assert key in data
        return response.json()
    return _validate