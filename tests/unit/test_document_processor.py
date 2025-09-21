"""
Unit tests for DocumentProcessor service.
"""

import pytest
import tempfile
import os
from unittest.mock import patch, Mock, mock_open
from io import BytesIO
import hashlib

from src.services.document_processor import DocumentProcessor


class TestDocumentProcessor:
    """Test cases for DocumentProcessor class."""

    @pytest.fixture
    def processor(self):
        """Create DocumentProcessor instance."""
        return DocumentProcessor()

    def test_calculate_file_hash(self):
        """Test file hash calculation."""
        content = b"test content"
        expected_hash = hashlib.sha256(content).hexdigest()
        actual_hash = DocumentProcessor.calculate_file_hash(content)
        assert actual_hash == expected_hash

    def test_detect_encoding_utf8(self):
        """Test encoding detection for UTF-8 content."""
        content = "Hello, world! 你好世界".encode('utf-8')
        encoding = DocumentProcessor.detect_encoding(content)
        assert encoding in ['utf-8', 'UTF-8']

    def test_detect_encoding_latin1(self):
        """Test encoding detection for Latin-1 content."""
        content = "Café résumé".encode('latin-1')
        encoding = DocumentProcessor.detect_encoding(content)
        assert encoding is not None

    def test_detect_encoding_fallback(self):
        """Test encoding detection fallback to utf-8."""
        with patch('chardet.detect', return_value={'encoding': None}):
            content = b"test content"
            encoding = DocumentProcessor.detect_encoding(content)
            assert encoding == 'utf-8'

    @pytest.mark.unit
    def test_validate_file_success(self, processor):
        """Test successful file validation."""
        content = b"Hello, world!"
        filename = "test.txt"

        with patch.object(processor, '_validate_file_signature', return_value=True):
            result = processor.validate_file(content, filename)

        assert result['valid'] is True
        assert result['file_type'] == 'txt'
        assert result['file_hash'] == hashlib.sha256(content).hexdigest()
        assert result['file_size'] == len(content)

    @pytest.mark.unit
    def test_validate_file_too_large(self, processor):
        """Test file validation with oversized file."""
        # Create content larger than max size
        large_content = b"x" * (50 * 1024 * 1024 + 1)  # 50MB + 1 byte
        filename = "large.txt"

        result = processor.validate_file(large_content, filename)

        assert result['valid'] is False
        assert "exceeds maximum" in result['error']

    @pytest.mark.unit
    def test_validate_file_unsupported_extension(self, processor):
        """Test file validation with unsupported file type."""
        content = b"test content"
        filename = "test.exe"

        result = processor.validate_file(content, filename)

        assert result['valid'] is False
        assert "not allowed" in result['error']

    @pytest.mark.unit
    def test_validate_file_signature_mismatch(self, processor):
        """Test file validation with signature mismatch."""
        content = b"not a pdf"
        filename = "test.pdf"

        result = processor.validate_file(content, filename)

        assert result['valid'] is False
        assert "does not match expected format" in result['error']

    def test_validate_file_signature_pdf_valid(self, processor):
        """Test PDF file signature validation."""
        content = b"%PDF-1.4\nPDF content"
        result = processor._validate_file_signature(content, 'pdf')
        assert result is True

    def test_validate_file_signature_pdf_invalid(self, processor):
        """Test invalid PDF file signature."""
        content = b"not a pdf"
        result = processor._validate_file_signature(content, 'pdf')
        assert result is False

    def test_validate_file_signature_docx_valid(self, processor):
        """Test DOCX file signature validation."""
        content = b"PK\x03\x04DOCX content"
        result = processor._validate_file_signature(content, 'docx')
        assert result is True

    def test_validate_file_signature_text_files(self, processor):
        """Test text file signatures (no validation needed)."""
        content = b"any content"
        assert processor._validate_file_signature(content, 'txt') is True
        assert processor._validate_file_signature(content, 'md') is True

    def test_validate_file_signature_empty_content(self, processor):
        """Test file signature validation with empty content."""
        content = b""
        result = processor._validate_file_signature(content, 'pdf')
        assert result is False

    @pytest.mark.unit
    def test_extract_text_unsupported_type(self, processor):
        """Test extract_text with unsupported file type."""
        with pytest.raises(ValueError, match="Unsupported file type: xyz"):
            processor.extract_text("dummy_path", "xyz")

    @pytest.mark.unit
    def test_extract_txt_text_success(self, processor, sample_txt_file, sample_text_content):
        """Test successful text file extraction."""
        result = processor.extract_text(sample_txt_file, "txt")

        assert result['text'] == sample_text_content
        assert result['total_pages'] == 1
        assert len(result['pages']) == 1
        assert result['pages'][0]['page_number'] == 1
        assert result['pages'][0]['text'] == sample_text_content
        assert 'encoding' in result['metadata']

    @pytest.mark.unit
    def test_extract_md_text_success(self, processor, sample_md_file, sample_text_content):
        """Test successful markdown file extraction."""
        result = processor.extract_text(sample_md_file, "md")

        assert result['text'] == sample_text_content
        assert result['total_pages'] == 1
        assert len(result['pages']) == 1
        assert result['metadata']['format'] == 'markdown'

    @pytest.mark.unit
    def test_extract_docx_text_success(self, processor, sample_docx_file):
        """Test successful DOCX file extraction."""
        result = processor.extract_text(sample_docx_file, "docx")

        assert result['text'] is not None
        assert len(result['text']) > 0
        assert result['total_pages'] == 1
        assert len(result['pages']) == 1
        assert result['pages'][0]['page_number'] == 1

    @pytest.mark.unit
    def test_extract_txt_text_encoding_error(self, processor, temp_dir):
        """Test text extraction with encoding issues."""
        # Create file with mixed encoding
        txt_path = os.path.join(temp_dir, "bad_encoding.txt")
        with open(txt_path, "wb") as f:
            f.write(b"Good text\xFF\xFEBad bytes")

        result = processor.extract_text(txt_path, "txt")

        # Should not fail due to error handling
        assert result['text'] is not None
        assert result['total_pages'] == 1

    @pytest.mark.unit
    def test_extract_txt_text_file_not_found(self, processor):
        """Test text extraction with non-existent file."""
        with pytest.raises(FileNotFoundError):
            processor.extract_text("/nonexistent/file.txt", "txt")

    @pytest.mark.unit
    def test_extract_docx_text_invalid_file(self, processor, temp_dir):
        """Test DOCX extraction with invalid file."""
        # Create a fake DOCX file
        docx_path = os.path.join(temp_dir, "fake.docx")
        with open(docx_path, "w") as f:
            f.write("not a docx")

        with pytest.raises(Exception):
            processor.extract_text(docx_path, "docx")

    @pytest.mark.unit
    def test_extract_docx_text_with_tables(self, processor, temp_dir):
        """Test DOCX extraction with tables."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Regular paragraph")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"

        docx_path = os.path.join(temp_dir, "with_table.docx")
        doc.save(docx_path)

        result = processor.extract_text(docx_path, "docx")

        assert "Regular paragraph" in result['text']
        assert "Cell 1" in result['text']
        assert "Cell 4" in result['text']

    @pytest.mark.unit
    @patch('pdfplumber.open')
    def test_extract_pdf_text_success(self, mock_pdfplumber, processor, temp_dir):
        """Test successful PDF extraction with pdfplumber."""
        # Mock PDF structure
        mock_page = Mock()
        mock_page.extract_text.return_value = "Page 1 content"

        mock_pdf = Mock()
        mock_pdf.pages = [mock_page]
        mock_pdf.metadata = {
            'Title': 'Test PDF',
            'Author': 'Test Author',
            'Subject': 'Test Subject'
        }

        mock_pdfplumber.return_value.__enter__.return_value = mock_pdf

        pdf_path = os.path.join(temp_dir, "test.pdf")

        result = processor.extract_text(pdf_path, "pdf")

        assert result['text'] == "Page 1 content"
        assert result['total_pages'] == 1
        assert len(result['pages']) == 1
        assert result['pages'][0]['text'] == "Page 1 content"
        assert result['metadata']['title'] == 'Test PDF'
        assert result['metadata']['author'] == 'Test Author'

    @pytest.mark.unit
    @patch('pdfplumber.open')
    def test_extract_pdf_text_multiple_pages(self, mock_pdfplumber, processor, temp_dir):
        """Test PDF extraction with multiple pages."""
        # Mock multiple pages
        mock_page1 = Mock()
        mock_page1.extract_text.return_value = "Page 1 content"

        mock_page2 = Mock()
        mock_page2.extract_text.return_value = "Page 2 content"

        mock_pdf = Mock()
        mock_pdf.pages = [mock_page1, mock_page2]
        mock_pdf.metadata = {}

        mock_pdfplumber.return_value.__enter__.return_value = mock_pdf

        pdf_path = os.path.join(temp_dir, "multipage.pdf")

        result = processor.extract_text(pdf_path, "pdf")

        assert result['total_pages'] == 2
        assert len(result['pages']) == 2
        assert "Page 1 content" in result['text']
        assert "Page 2 content" in result['text']

    @pytest.mark.unit
    @patch('pdfplumber.open')
    def test_extract_pdf_text_page_error(self, mock_pdfplumber, processor, temp_dir):
        """Test PDF extraction with page extraction error."""
        # Mock page that raises exception
        mock_page = Mock()
        mock_page.extract_text.side_effect = Exception("Page extraction failed")

        mock_pdf = Mock()
        mock_pdf.pages = [mock_page]
        mock_pdf.metadata = {}

        mock_pdfplumber.return_value.__enter__.return_value = mock_pdf

        pdf_path = os.path.join(temp_dir, "error.pdf")

        result = processor.extract_text(pdf_path, "pdf")

        assert result['total_pages'] == 1
        assert result['pages'][0]['text'] == ""
        assert result['pages'][0]['char_count'] == 0

    @pytest.mark.unit
    @patch('pdfplumber.open')
    @patch('PyPDF2.PdfReader')
    def test_extract_pdf_text_fallback_to_pypdf2(self, mock_pypdf2, mock_pdfplumber, processor, temp_dir):
        """Test PDF extraction fallback to PyPDF2."""
        # Make pdfplumber fail
        mock_pdfplumber.side_effect = Exception("pdfplumber failed")

        # Mock PyPDF2
        mock_page = Mock()
        mock_page.extract_text.return_value = "PyPDF2 content"

        mock_reader = Mock()
        mock_reader.pages = [mock_page]
        mock_reader.metadata = {'/Title': 'PDF Title'}

        mock_pypdf2.return_value = mock_reader

        pdf_path = os.path.join(temp_dir, "fallback.pdf")

        result = processor.extract_text(pdf_path, "pdf")

        assert result['text'] == "PyPDF2 content"
        assert result['metadata']['title'] == 'PDF Title'

    @pytest.mark.unit
    @patch('pdfplumber.open')
    @patch('PyPDF2.PdfReader')
    def test_extract_pdf_text_complete_failure(self, mock_pypdf2, mock_pdfplumber, processor, temp_dir):
        """Test PDF extraction when both libraries fail."""
        # Make both libraries fail
        mock_pdfplumber.side_effect = Exception("pdfplumber failed")
        mock_pypdf2.side_effect = Exception("PyPDF2 failed")

        pdf_path = os.path.join(temp_dir, "failed.pdf")

        with pytest.raises(Exception):
            processor.extract_text(pdf_path, "pdf")

    def test_supported_formats_constant(self, processor):
        """Test that supported formats are properly defined."""
        formats = processor.SUPPORTED_FORMATS

        assert 'pdf' in formats
        assert 'docx' in formats
        assert 'txt' in formats
        assert 'md' in formats

        assert 'application/pdf' in formats['pdf']
        assert 'text/plain' in formats['txt']
        assert 'text/markdown' in formats['md']